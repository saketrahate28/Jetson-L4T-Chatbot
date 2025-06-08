# Standard library imports
import os
import threading
import logging
from flask import Flask, request, jsonify, send_from_directory, Markup
from flask_cors import CORS
import openai
import docx
import pandas as pd
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_nvidia_ai_endpoints import ChatNVIDIA, NVIDIAEmbeddings
from langchain_core.documents import Document
import re
import time
from dotenv import load_dotenv
import markdown2 as markdown
import bleach

load_dotenv()

NVIDIA_API_KEY = os.getenv("NVIDIA_API_KEY")

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Clear SSL issues
if 'SSL_CERT_FILE' in os.environ:
    os.environ.pop('SSL_CERT_FILE')

os.environ["NVIDIA_API_KEY"] = NVIDIA_API_KEY
print(f"NVIDIA API Key configured: {NVIDIA_API_KEY[:10] if NVIDIA_API_KEY else 'Not found'}...")

base_url = "https://integrate.api.nvidia.com/v1"
model="nvdev/meta/llama-3.1-70b-instruct"
max_tokens=1024
temperature=0.5

app = Flask(__name__, static_folder="../UI")
CORS(app, resources={r"/*": {"origins": "*"}}, supports_credentials=True)

# Configure OpenAI API
openai.api_key = NVIDIA_API_KEY
openai.api_base = base_url

# Track embedding status
embeddings_ready = False
vectors = None
embeddings = None
docstore = None
convstore = None

# Add this near the top of the file, with other globals
response_cache = {}
CACHE_MAX_SIZE = 100  # Maximum number of cached responses to keep

embedder = NVIDIAEmbeddings(model="nvidia/nv-embed-v1")
instruct_llm = ChatNVIDIA(model=model, max_tokens=max_tokens, temperature=temperature)

def docs2str(docs, title="Document"):
    """Useful utility for making chunks into context string."""
    out_str = ""
    for doc in docs:
        doc_name = getattr(doc, 'metadata', {}).get('Title', title)
        if doc_name:
            out_str += f"[Quote from {doc_name}] "
        out_str += getattr(doc, 'page_content', str(doc)) + "\n"
    return out_str

    # Function to save conversation history
def save_memory_and_get_output(d, vstore):
    """Accepts 'input'/'output' dictionary and saves to conversation store"""
    vstore.add_texts([f"User said {d.get('input')}", f"Agent said {d.get('output')}"])
    return d.get('output')

def initialize_embeddings():
    global embeddings_ready, vectors, embeddings, docstore, convstore

    # Only initialize once
    if embeddings_ready and docstore is not None:
        print("Embeddings already initialized, skipping...")
        return

    print("Starting document embeddings initialization...")
    try:
        # Make sure Data directory exists
        data_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Data')
        print(f"Looking for Data folder at: {data_path}")

        # Create the Data directory if it doesn't exist
        if not os.path.exists(data_path):
            print(f"Data directory not found. Creating it at: {data_path}")
            os.makedirs(data_path, exist_ok=True)
            print("Data directory created. Please add documents to this folder.")
            embeddings_ready = False
            return

        # Create an empty list to store all documents
        documents = []
        processed_files = set()  # Track files we've already processed

        # Process each file individually by type
        for root, _, files in os.walk(data_path):
            for file in files:
                file_path = os.path.join(root, file)

                # Skip if we already processed this file
                if file_path in processed_files:
                    continue

                processed_files.add(file_path)

                # Skip JavaScript files and other non-document types
                if file.endswith('.js'):
                    continue

                try:
                    # Use appropriate loader based on file extension
                    file_lower = file.lower()

                    if file_lower.endswith('.pdf'):
                        loader = PyPDFLoader(file_path)
                        documents.extend(loader.load())
                        print(f"Loaded PDF: {file}")

                    elif file_lower.endswith('.docx'):
                        # Process DOCX files
                        if docx:  # Check if docx module is available
                            doc = docx.Document(file_path)
                            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

                            # Also extract tables if present
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = " | ".join([cell.text for cell in row.cells])
                                    text_content += "\n" + row_text

                            documents.append(Document(
                                page_content=text_content,
                                metadata={"source": file_path, "title": file}
                            ))
                            print(f"Loaded DOCX: {file}")
                        else:
                            print(f"Error loading {file}: Missing python-docx module")
                            print("Install with: pip install python-docx")

                    elif file_lower.endswith('.txt'):
                        loader = TextLoader(file_path)
                        documents.extend(loader.load())
                        print(f"Loaded TXT: {file}")

                    elif file_lower.endswith('.xlsx') or file_lower.endswith('.xls'):
                        # Process Excel files
                        try:
                            excel_file = pd.ExcelFile(file_path)
                            for sheet_name in excel_file.sheet_names:
                                df = pd.read_excel(file_path, sheet_name=sheet_name)
                                # Convert DataFrame to string representation
                                text_content = f"Excel File: {file}, Sheet: {sheet_name}\n\n"

                                # Add column headers
                                text_content += "| " + " | ".join(str(col) for col in df.columns) + " |\n"
                                text_content += "| " + " | ".join("-" * len(str(col)) for col in df.columns) + " |\n"

                                # Add rows
                                for _, row in df.iterrows():
                                    text_content += "| " + " | ".join(str(cell) for cell in row) + " |\n"

                                # Create document
                                doc = Document(
                                    page_content=text_content,
                                    metadata={"source": file_path, "sheet": sheet_name}
                                )
                                documents.append(doc)
                            print(f"Loaded Excel: {file}")
                        except Exception as excel_error:
                            print(f"Error processing Excel file {file}: {str(excel_error)}")
                    else:
                        print(f"Skipped unsupported file: {file}")

                except Exception as e:
                    print(f"Error loading {file}: {str(e)}")
                    continue

        if documents:
            print(f"Number of Documents Retrieved: {len(documents)}")

            # Split documents
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1200,
                chunk_overlap=100,
                separators=["\n\n", "\n", ".", ";", ",", " ", ""]
            )

            docs_split = text_splitter.split_documents(documents)
            print(f"Total chunks after splitting: {len(docs_split)}")

            # Create FAISS index from document chunks - with optimized batching
            # Calculate optimal batch size based on document count
            total_docs = len(docs_split)

            # Dynamic batch sizing: smaller batches for large collections,
            # single batch for small collections
            if total_docs <= 100:
                optimal_batch = total_docs  # Process all at once if ≤ 100 documents
            else:
                # For larger collections, scale batch size to avoid overwhelming API
                # Larger collections get proportionally smaller batches
                optimal_batch = max(25, min(100, total_docs // 10))

            print(f"Creating document store with optimal batch size: {optimal_batch} for {total_docs} documents")

            # If there are too many documents, process in batches
            if total_docs > optimal_batch:
                all_batches = [docs_split[i:i + optimal_batch] for i in range(0, total_docs, optimal_batch)]
                print(f"Processing {len(all_batches)} batches")

                # Process first batch to initialize the index
                print(f"Processing initial batch of {len(all_batches[0])} documents")
                docstore = FAISS.from_documents(all_batches[0], embedder)

                # Process remaining batches
                for i, batch in enumerate(all_batches[1:], 1):
                    print(f"Processing batch {i+1}/{len(all_batches)} with {len(batch)} documents")
                    # Add documents to existing index
                    docstore.add_documents(batch)
            else:
                # Process all at once if small enough
                print(f"Processing all {total_docs} documents at once")
                docstore = FAISS.from_documents(docs_split, embedder)

            print(f"Created document store with {total_docs} chunks")

            # Initialize conversation history
            conversation = [
                "User asked about document analysis",
                "Agent explained how documents are processed",
                "User asked about embeddings",
                "Agent explained vector embeddings for text"
            ]

            # Create conversation vector store
            convstore = FAISS.from_texts(conversation, embedding=embedder)
            print(f"Created conversation store with {len(conversation)} entries")

            embeddings_ready = True
            print("Document embeddings initialization completed successfully")

        else:
            print("No documents were loaded from the directory")

    except Exception as e:
        print(f"Error initializing embeddings: {str(e)}")
        import traceback
        traceback.print_exc()
        return

# Start embeddings initialization in a background thread
threading.Thread(target=initialize_embeddings).start()

# Routes
@app.route('/')
def index():
    return send_from_directory('../UI', 'chatbot.html')

@app.route('/test-rag')
def test_rag():
    """Serve the test-rag.html interface"""
    return send_from_directory('../UI', 'test-rag.html')

@app.route('/api/status', methods=['GET'])
def status():
    return jsonify({
        'nim_api': 'configured',
        'embeddings_ready': embeddings_ready
    })

@app.route('/api/query-documents', methods=['POST'])
def query_documents():
    global docstore, embeddings, embeddings_ready

    data = request.json
    query = data.get('query')

    if not query:
        return jsonify({
            'status': 'error',
            'message': 'Query is required'
        }), 400

    if not embeddings_ready:
        return jsonify({
            'status': 'error',
            'message': 'Document embeddings are not ready yet. Please try again later.'
        }), 503

    try:
        # Perform search
        results = docstore.similarity_search_with_score(query, k=5)

        # Format results
        formatted_results = []
        for doc, score in results:
            formatted_results.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": float(score)
            })

        return jsonify({
            'status': 'success',
            'data': {
                'results': formatted_results
            }
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

# Function to verify API key
def verify_api_key(request):
    """Verify the API key in the request headers"""
    auth_header = request.headers.get('Authorization')
    if not auth_header:
        return False

    # Extract the token part after "Bearer "
    try:
        parts = auth_header.split()
        if parts[0].lower() != 'bearer':
            return False
        token = parts[1]
        # Compare with our configured API key
        return token == NVIDIA_API_KEY
    except:
        return False

@app.route('/api/chat', methods=['POST'])
def chat():
    """Process chat messages with NVIDIA API"""
    try:
        print("Received chat request")
        data = request.json
        if not data or not isinstance(data.get('messages'), list):
            return jsonify({
                "status": "error",
                "message": "Invalid request: messages array is required"
            }), 400

        messages = data.get('messages', [])

        # Make API request using older OpenAI API style
        response = openai.ChatCompletion.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
            request_timeout=30  # Add 30 second timeout to prevent hanging
        )

        response_text = response.choices[0].message.content

        return jsonify({
            "status": "success",
            "message": response_text
        })
    except Exception as e:
        print(f"Error in chat endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/rag_chat', methods=['POST'])
def rag_chat():
    """RAG-enhanced chat endpoint"""
    global docstore, convstore, embeddings_ready, response_cache

    print("RAG chat endpoint called")

    # Add detailed request logging
    print("Request headers:", dict(request.headers))
    print("Request body:", request.get_data(as_text=True))

    try:
        data = request.json
        if not data:
            return jsonify({
                "status": "error",
                "message": "Invalid request: empty request body"
            }), 400

        query = data.get('query', '')

        print(f"RAG query: {query}")
        print(f"Embeddings ready: {embeddings_ready}")
        print(f"Docstore exists: {docstore is not None}")

        if not query:
            return jsonify({
                "status": "error",
                "message": "Query parameter is required"
            }), 400

        # Check cache for exact match
        cache_key = query.strip().lower()
        if cache_key in response_cache:
            print(f"Cache hit for query: {query[:30]}...")
            return jsonify(response_cache[cache_key])

        # Check if embeddings are ready
        if not embeddings_ready or docstore is None:
            print("WARNING: Document embeddings not ready yet, falling back to basic chat")

            # Fall back to basic chat if embeddings aren't ready
            try:
                # Use direct OpenAI call instead of langchain
                response = openai.ChatCompletion.create(
                    model=model,
                    messages=[{"role": "user", "content": query}],
                    temperature=temperature,
                    max_tokens=1024,
                    request_timeout=30  # Add 30 second timeout to prevent hanging
                )
                response_content = response.choices[0].message.content

                result = {
                    "status": "success",
                    "message": response_content,
                    "debug_info": {
                        "used_rag": False,
                        "reason": "Embeddings not ready",
                        "embeddings_ready": embeddings_ready,
                        "docstore_exists": docstore is not None
                    }
                }

                # Cache the response
                if len(response_cache) >= CACHE_MAX_SIZE:
                    # Remove a random key if cache is full
                    random_key = next(iter(response_cache))
                    response_cache.pop(random_key)

                response_cache[cache_key] = result
                return jsonify(result)
            except Exception as e:
                print(f"Error in fallback chat: {str(e)}")
                return jsonify({
                    "status": "error",
                    "message": f"Error in fallback chat: {str(e)}",
                    "debug_info": {
                        "used_rag": False,
                        "reason": "Embeddings not ready + fallback failed"
                    }
                }), 500

        try:
            # SIMPLIFIED RAG APPROACH - bypass complex LangChain chains
            # 1. Get relevant documents
            print("Retrieving relevant documents")
            results = docstore.similarity_search_with_score(query, k=5)

            if not results:
                print("No relevant documents found, falling back to regular chat")
                response = openai.ChatCompletion.create(
                    model=model,
                    messages=[{"role": "user", "content": query}],
                    temperature=temperature,
                    max_tokens=1024,
                    request_timeout=30
                )
                response_content = response.choices[0].message.content
                result = {
                    "status": "success",
                    "message": response_content,
                    "debug_info": {
                        "used_rag": False,
                        "reason": "No relevant documents found"
                    }
                }
                if len(response_cache) >= CACHE_MAX_SIZE:
                    random_key = next(iter(response_cache))
                    response_cache.pop(random_key)
                response_cache[cache_key] = result
                return jsonify(result)

            # 2. Format context from documents
            context = ""
            for doc, score in results:
                if score > 0.75:
                    source_path = doc.metadata.get('source', 'Unknown')
                    source_filename = source_path.split('/')[-1].split('\\')[-1]
                    context += f"\n[From {source_filename}]: {doc.page_content}\n"

            print(f"Retrieved context (first 100 chars): {context[:100]}...")

            # 3. Create prompt with context
            prompt = f"""Answer the following question using clean, modern Markdown.
If your answer includes commands, always include the full commands inside properly formatted Markdown code blocks (using triple backticks). Do not leave code blocks empty. Do not omit any commands.
Structure your answer with clear sections and headings (#, ##, ###), use bullet points (-) and numbered lists (1., 2., 3.) where appropriate, and write in detailed, readable paragraphs. Explain in detail wherever needed, don't just give key points. Highlight important terms with **bold**. Do not underline each section before and after. Do not use ASCII art, manual underlines, equals signs, or visual dividers. Do not show Markdown syntax in the output—just use proper Markdown for structure. Avoid excessive punctuation, repeated lines, or asterisks for lists. Only use code blocks if specifically requested. No citations or source references.
Do not add lines after and before the sections, eliminate the unnecessary white spaces between the sections.

DOCUMENTS:
{context}

QUESTION: {query}

ANSWER:"""

            print("Sending to LLM")
            start_time = time.time()
            print(f"Created LLM prompt with {len(prompt)} characters, sending to API...")

            try:
                response = openai.ChatCompletion.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=temperature,
                    max_tokens=1024,
                    request_timeout=45
                )
                response_content = response.choices[0].message.content
                # Minimal cleanup: normalize line endings and double spaces
                response_content = re.sub(r'\r\n', '\n', response_content)
                response_content = re.sub(r' +', ' ', response_content)
                # Remove citation patterns (privacy)
                citation_pattern = r'\[From [^\]]*\]'
                response_content = re.sub(citation_pattern, '', response_content)
                # Remove source references (privacy)
                source_reference_pattern = r'(?:According to|From|In|As mentioned in|As stated in|As described in|As shown in|As per|Based on)[^.,:;]*\.[a-zA-Z0-9]+'
                response_content = re.sub(source_reference_pattern, '', response_content)
                # Clean up unnecessary white spaces between sections
                response_content = re.sub(r'\n{3,}', '\n\n', response_content)  # Replace 3+ newlines with 2
                response_content = re.sub(r'(\n#+\s.*)\n+', r'\1\n', response_content)  # Remove extra newlines after headers
                response_content = re.sub(r'\n+(\n#+\s.*)', r'\n\1', response_content)  # Remove extra newlines before headers
                response_content = re.sub(r'(\n- .*)\n+', r'\1\n', response_content)  # Remove extra newlines after list items
                response_content = re.sub(r'\n+(\n- .*)', r'\n\1', response_content)  # Remove extra newlines before list items
                response_content = re.sub(r'(\n\d+\. .*)\n+', r'\1\n', response_content)  # Remove extra newlines after numbered list items
                response_content = re.sub(r'\n+(\n\d+\. .*)', r'\n\1', response_content)  # Remove extra newlines before numbered list items
                result = {
                    "status": "success",
                    "message": response_content,
                    "format": "markdown",
                    "debug_info": {
                        "used_rag": True,
                        "context_length": len(context),
                        "retrieved_docs": len(results),
                        "llm_time_seconds": time.time() - start_time
                    }
                }
                if len(response_cache) >= CACHE_MAX_SIZE:
                    random_key = next(iter(response_cache))
                    response_cache.pop(random_key)
                response_cache[cache_key] = result
                return jsonify(result)
            except Exception as e:
                print(f"Error in RAG processing: {str(e)}")
                import traceback
                traceback.print_exc()

        except Exception as e:
            print(f"Error in RAG processing: {str(e)}")
            import traceback
            traceback.print_exc()

            # Fall back to regular chat if RAG fails
            try:
                # Use direct OpenAI call instead of langchain
                fallback_response = openai.ChatCompletion.create(
                    model=model,
                    messages=[{"role": "user", "content": query}],
                    temperature=temperature,
                    max_tokens=1024,
                    request_timeout=30  # Add 30 second timeout to prevent hanging
                )
                fallback_content = fallback_response.choices[0].message.content

                result = {
                    "status": "fallback",
                    "message": fallback_content,
                    "debug_info": {
                        "rag_error": str(e),
                        "used_fallback": True
                    }
                }

                # Cache the response
                if len(response_cache) >= CACHE_MAX_SIZE:
                    # Remove a random key if cache is full
                    random_key = next(iter(response_cache))
                    response_cache.pop(random_key)

                response_cache[cache_key] = result
                return jsonify(result)
            except Exception as fallback_error:
                print(f"Fallback also failed: {str(fallback_error)}")
                return jsonify({
                    "status": "error",
                    "message": f"RAG processing failed and fallback also failed. Primary error: {str(e)}",
                    "debug_info": {
                        "rag_error": str(e),
                        "fallback_error": str(fallback_error)
                    }
                }), 500

    except Exception as e:
        print(f"Error in rag_chat endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}",
            "error": str(e)
        }), 500

@app.route('/api/generator', methods=['POST'])
def generator():
    """Content generation endpoint"""
    try:
        data = request.json
        prompt = data.get('prompt', '')

        if not prompt:
            return jsonify({
                "status": "error",
                "message": "Prompt parameter is required"
            }), 400

        # Create generation prompt
        generation_prompt = ChatPromptTemplate.from_template(
            "Generate creative content based on the following prompt: {prompt}\n\n"
            "Be detailed, engaging, and original in your response."
        )

        # Create generation chain
        generation_chain = (
            {'prompt': (lambda x: x)}
            | generation_prompt
            | instruct_llm
            | StrOutputParser()
        )

        # Get response
        response_text = generation_chain.invoke(prompt)

        return jsonify({
            "status": "success",
            "message": response_text
        })
    except Exception as e:
        print(f"Error in generator endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/retriever', methods=['POST'])
def retriever():
    """Document retrieval endpoint"""
    global docstore, embeddings_ready

    try:
        data = request.json
        query = data.get('query', '')

        if not query:
            return jsonify({
                "status": "error",
                "message": "Query parameter is required"
            }), 400

        if not embeddings_ready:
            return jsonify({
                "status": "error",
                "message": "Document embeddings are not ready yet. Please try again later."
            }), 503

        # Perform search
        results = docstore.similarity_search_with_score(query, k=5)

        # Format results
        formatted_results = []
        for doc, score in results:
            formatted_results.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": float(score)
            })

        return jsonify({
            "status": "success",
            "data": formatted_results
        })
    except Exception as e:
        print(f"Error in retriever endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/embeddings', methods=['POST'])
def create_embeddings():
    """Endpoint to create embeddings for texts"""
    try:
        data = request.json
        texts = data.get('texts', [])

        if not texts:
            return jsonify({
                "status": "error",
                "message": "No texts provided"
            }), 400

        # Use NVIDIAEmbeddings to create embeddings
        embeddings = embedder.embed_documents(texts)

        return jsonify({
            "status": "success",
            "embeddings": embeddings
        })
    except Exception as e:
        print(f"Error in embeddings endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/<path:filename>')
def serve_static(filename):
    # Try UI folder first
    ui_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'UI')
    if os.path.exists(os.path.join(ui_path, filename)):
        return send_from_directory('../UI', filename)

    # Then try Data folder
    data_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Data')
    if os.path.exists(os.path.join(data_path, filename)):
        return send_from_directory('../Data', filename)

    return "File not found", 404

@app.route('/api/nim/status', methods=['GET'])
def nim_status():
    """Endpoint to check NVIDIA NIM API status"""
    try:
        # Test the connection with a minimal query using OpenAI API style
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": "Hi"}],
            max_tokens=5,
            temperature=0.1,
            request_timeout=10  # Add 10 second timeout to prevent hanging
        )

        # Connection successful
        return jsonify({
            "connected": True,
            "status": "ok",
            "model": model,
            "base_url": base_url,
            "embeddings_ready": embeddings_ready
        })
    except Exception as e:
        print(f"Error checking NIM status: {str(e)}")
        return jsonify({
            "connected": False,
            "status": "error",
            "error": str(e)
        }), 500

@app.route('/api/debug', methods=['GET'])
def debug_info():
    """Diagnostic endpoint to check system status"""
    global embeddings_ready, docstore, convstore

    # Check if embeddings are actually initialized
    embedding_status = {
        "embeddings_ready_flag": embeddings_ready,
        "docstore_exists": docstore is not None,
        "convstore_exists": convstore is not None,
    }

    # Get counts of documents if possible
    if docstore is not None:
        try:
            embedding_status["docstore_index_ntotal"] = docstore.index.ntotal
        except:
            embedding_status["docstore_index_ntotal"] = "error accessing"

    # Testing retrieval on a simple query if docstore exists
    retrieval_test = {}
    if docstore is not None:
        try:
            results = docstore.similarity_search_with_score("Jetson", k=1)
            retrieval_test["success"] = True
            retrieval_test["result_count"] = len(results)
            if results:
                doc, score = results[0]
                retrieval_test["first_result"] = {
                    "content_preview": doc.page_content[:100] + "...",
                    "score": float(score),
                    "metadata": doc.metadata
                }
        except Exception as e:
            retrieval_test["success"] = False
            retrieval_test["error"] = str(e)

    # Testing the LLM directly
    llm_test = {}
    try:
        # Use direct OpenAI call instead of langchain
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": "Say hello"}],
            temperature=temperature,
            max_tokens=20,
            request_timeout=10  # Add 10 second timeout to prevent hanging
        )
        llm_test["success"] = True
        llm_test["response"] = response.choices[0].message.content
    except Exception as e:
        llm_test["success"] = False
        llm_test["error"] = str(e)

    return jsonify({
        "server_status": "running",
        "embeddings": embedding_status,
        "retrieval_test": retrieval_test,
        "llm_test": llm_test,
        "api_key_configured": bool(NVIDIA_API_KEY),
        "model": model
    })

@app.route('/api/get_api_key', methods=['GET'])
def get_api_key():
    """Return a masked version of the API key"""
    if NVIDIA_API_KEY:
        masked_key = NVIDIA_API_KEY[:4] + "..." + NVIDIA_API_KEY[-4:]
        return jsonify({
            "status": "success",
            "key_configured": True,
            "masked_key": masked_key
        })
    else:
        return jsonify({
            "status": "error",
            "key_configured": False,
            "message": "No API key configured on the server"
        })

@app.route('/nim-visualization')
def nim_visualization():
    """Serve the NIM visualization page"""
    return send_from_directory('../UI', 'nim-visualization.html')

if __name__ == '__main__':
    print("Starting Flask server...")
    app.run(debug=True, threaded=True, port=3000)